from django.shortcuts import render
from .models import SoilData
from .models import Predicted
from django.http import HttpResponse
from django.template.loader import render_to_string
from reportlab.pdfgen import canvas
from io import BytesIO
import json
from django.core.serializers.json import DjangoJSONEncoder
from docx import Document

def admin_dashboard(request):
    latest_reading = SoilData.objects.latest('timestamp')
    readings = SoilData.objects.order_by('-timestamp')[:100]
    readings_json = json.dumps(list(readings.values()), cls=DjangoJSONEncoder)
    
    context = {
        'latest_reading': latest_reading,
        'readings_json': readings_json,
    }
    
    return render(request, 'adminApp/dashboard.html', context)

# def prediction_table(request):
#     try:
#         # Fetch all soil data ordered by timestamp (most recent first)
#         soil_data = SoilData.objects.all().order_by('-timestamp')
        
#         # Get the latest reading for the stats cards
#         latest_reading = SoilData.objects.latest('timestamp')
        
#         context = {
#             'measurements': soil_data,
#             'latest_reading': latest_reading,
#         }
        
#         return render(request, 'adminApp/prediction_table.html', context)
        
#     except Exception as e:
#         print(f"Error fetching soil data: {str(e)}")
#         context = {
#             'measurements': [],
#             'latest_reading': None,
#             'error_message': 'Error fetching soil data'
#         }
#         return render(request, 'adminApp/prediction_table.html', context)

def prediction_table(request):
    try:
        # Fetch all soil data ordered by timestamp (most recent first)
        soil_data = SoilData.objects.all().order_by('-timestamp')
        
        # Initialize a list to hold the highest-rated crop data for each soil entry
        prediction_data = []

        for measurement in soil_data:
            # Get the predicted crop entry for the current soil data
            predicted_crop_entry = Predicted.objects.filter(soil_data_id=measurement.id).first()
            
            if predicted_crop_entry and predicted_crop_entry.predicted_crops:
                # Extract the list of predicted crops
                predicted_crops = predicted_crop_entry.predicted_crops
                
                # Find the crop with the highest success rate
                highest_rate_crop = max(
                    predicted_crops,
                    key=lambda crop: float(crop['success_rate'].strip('%'))
                )
                
                # Add the relevant data to the prediction_data list
                prediction_data.append({
                    'soil_data_id': measurement.id,
                    'highest_rate_crop': highest_rate_crop,
                    'prediction_datetime': predicted_crop_entry.prediction_datetime,
                })
        
        context = {
            'measurements': soil_data,
            'latest_reading': soil_data.first(),  # Assuming latest_reading is the first entry in the list
            'prediction_data': prediction_data,   # Pass the collected prediction data
        }
        
        return render(request, 'adminApp/prediction_table.html', context)
        
    except Exception as e:
        print(f"Error fetching soil data: {str(e)}")
        context = {
            'measurements': [],
            'latest_reading': None,
            'prediction_data': [],
            'error_message': 'Error fetching soil data'
        }
        return render(request, 'adminApp/prediction_table.html', context)


def export_measurements(request):
    if request.method == 'POST':
        format = request.POST.get('format', 'pdf')
        measurements = SoilData.objects.all().order_by('-timestamp')
        
        if format == 'pdf':
            buffer = BytesIO()
            p = canvas.Canvas(buffer)
            
            y = 800
            p.drawString(100, y, "Soil Analysis Measurements")
            y -= 20
            
            for measurement in measurements:
                p.drawString(100, y, f"Date: {measurement.timestamp}")
                p.drawString(100, y-15, f"Nitrogen: {measurement.nitrogen}")
                p.drawString(100, y-30, f"Phosphorus: {measurement.phosphorus}")
                p.drawString(100, y-45, f"Potassium: {measurement.potassium}")
                p.drawString(100, y-60, f"pH: {measurement.soilPH}")
                p.drawString(100, y-75, f"Temperature: {measurement.temperature}")
                p.drawString(100, y-90, f"Humidity: {measurement.humidity}")
                y -= 100
            
            p.showPage()
            p.save()
            
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="soil_measurements.pdf"'
            
            return response
            
        elif format == 'doc':
            doc = Document()
            doc.add_heading('Soil Analysis Measurements', 0)
            
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Date'
            header_cells[1].text = 'Nitrogen'
            header_cells[2].text = 'Phosphorus'
            header_cells[3].text = 'Potassium'
            header_cells[4].text = 'pH'
            header_cells[5].text = 'Temperature'
            header_cells[6].text = 'Humidity'
            
            for measurement in measurements:
                row_cells = table.add_row().cells
                row_cells[0].text = str(measurement.timestamp)
                row_cells[1].text = str(measurement.nitrogen)
                row_cells[2].text = str(measurement.phosphorus)
                row_cells[3].text = str(measurement.potassium)
                row_cells[4].text = str(measurement.soilPH)
                row_cells[5].text = str(measurement.temperature)
                row_cells[6].text = str(measurement.humidity)
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename="soil_measurements.docx"'
            
            return response
    
    return HttpResponse('Invalid request method')

